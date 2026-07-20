from django.http import JsonResponse, HttpResponse
from django.shortcuts import render
from django.views.decorators.csrf import csrf_exempt
import PyPDF2
import io
import os
import re
import json
from datetime import datetime
from .models import ChatSession, ChatMessage, SessionFile, SessionMemory
from google import genai


def index(request):
    return render(request, "chat/index.html")


# ─────────────────────────────────────────────────────────────
# Access Memory Endpoints
# ─────────────────────────────────────────────────────────────

def get_session_memory(request, session_id):
    try:
        memory = SessionMemory.objects.get(session_id=session_id)
        return JsonResponse({"content": memory.content})
    except SessionMemory.DoesNotExist:
        return JsonResponse({"content": ""})


@csrf_exempt
def save_session_memory(request, session_id):
    # Accept both JSON and form-encoded bodies
    content = ""
    if request.content_type and "application/json" in request.content_type:
        try:
            body = json.loads(request.body)
            content = body.get("content", "").strip()
        except json.JSONDecodeError:
            pass
    else:
        content = (request.POST.get("content") or "").strip()

    try:
        session = ChatSession.objects.get(id=session_id)
        memory, _ = SessionMemory.objects.get_or_create(session=session)
        memory.content = content
        memory.save()
        return JsonResponse({"success": True})
    except ChatSession.DoesNotExist:
        return JsonResponse({"error": "Session not found"}, status=404)


# ─────────────────────────────────────────────────────────────
# Helper: Manuals
# ─────────────────────────────────────────────────────────────

def load_all_manuals():
    manuals_dir = os.path.join(os.path.dirname(__file__), 'manuals')
    manuals = {}
    if not os.path.exists(manuals_dir):
        os.makedirs(manuals_dir)
    for filename in os.listdir(manuals_dir):
        if filename.endswith('.txt'):
            manual_name = filename.replace('.txt', '').replace('_', ' ').title()
            try:
                with open(os.path.join(manuals_dir, filename), 'r', encoding='utf-8') as f:
                    manuals[manual_name] = f.read()
            except Exception:
                continue
    if not manuals:
        manuals['General'] = "No manuals found. Using general knowledge."
    return manuals


# ─────────────────────────────────────────────────────────────
# Helper: Document Type Detection
# ─────────────────────────────────────────────────────────────

def detect_document_type(text, filename):
    text_lower = text.lower()
    filename_lower = filename.lower()
    if 'cv' in filename_lower or 'resume' in filename_lower:
        return "CV"
    cv_keywords = ['experience', 'education', 'skills', 'employment', 'work history']
    if sum(1 for kw in cv_keywords if kw in text_lower) >= 3:
        return "CV"
    return "DATA" if filename.endswith('.csv') else "DOCUMENT"


# ─────────────────────────────────────────────────────────────
# Helper: Full CV Scoring Rubric (max 115 pts)
# ─────────────────────────────────────────────────────────────

def calculate_scores_from_cv_texts(cv_texts):
    results = []

    for idx, text in enumerate(cv_texts):
        t = text.lower()

        # ── Candidate Name ──
        name = f"Candidate {idx + 1}"
        skip = {'curriculum vitae', 'cv', 'resume', 'résumé', 'profile',
                'personal information', 'personal details', 'contact', 'about me'}
        for line in text.splitlines():
            line = line.strip()
            if (line and len(line.split()) <= 5
                    and not any(c.isdigit() for c in line)
                    and line.lower() not in skip
                    and not all(c in '=-_*#~+' for c in line.replace(' ', ''))):
                name = line
                break

        # ── Education (max 30 pts) ──
        edu_points = 0
        gpa_label = "No GPA found"
        gpa_match = re.search(r'gpa[:\s]*([0-9]\.[0-9]+)', t)
        if gpa_match:
            gpa = float(gpa_match.group(1))
            gpa_label = f"GPA {gpa:.2f}"
            if gpa >= 2.5:
                edu_points = 30
            elif gpa >= 1.5:
                edu_points = 10
            else:
                edu_points = 0

        # ── Work Experience (max 45 pts) ──
        exp_points = 0
        exp_label = "No experience found"
        year_matches = re.findall(r'(\d+(?:\.\d+)?)\s*(?:\+\s*)?years?', t)
        month_matches = re.findall(r'(\d+)\s*months?', t)
        total_years = 0.0
        if year_matches:
            total_years = max(float(y) for y in year_matches)
        elif month_matches:
            total_years = max(int(m) for m in month_matches) / 12

        if total_years >= 4.0:
            exp_points = 45
            exp_label = f"{total_years:.1f}+ years"
        elif total_years >= 1.0:
            exp_points = 30
            exp_label = f"{total_years:.1f} years"
        elif total_years > 0:
            months = int(total_years * 12)
            if months >= 3:
                exp_points = 15
                exp_label = f"{months} months"
            else:
                exp_points = 0
                exp_label = f"{months} months (< 3 months)"

        # ── Technical Skills (max 20 pts) ──
        tech_points = 0
        tech_label = "None detected"
        tier_a = ['python', 'c#', 'c++', 'sql']
        tier_b = ['java', 'html', 'css', 'javascript', 'node.js', 'nodejs']
        has_a = any(skill in t for skill in tier_a)
        has_b = any(skill in t for skill in tier_b)
        if has_a:
            found = [s for s in tier_a if s in t]
            tech_points = 20
            tech_label = f"Tier A ({', '.join(found)})"
        elif has_b:
            found = [s for s in tier_b if s in t]
            tech_points = 5
            tech_label = f"Tier B ({', '.join(found)})"

        # ── Projects (max 20 pts) ──
        proj_points = 0
        proj_label = "0 projects"
        proj_keyword_count = len(re.findall(r'\bproject\b', t))
        proj_count = min(proj_keyword_count, 10)
        if proj_count >= 3:
            proj_points = 20
            proj_label = f"{proj_count} projects"
        elif proj_count >= 1:
            proj_points = 10
            proj_label = f"{proj_count} project(s)"

        total = edu_points + exp_points + tech_points + proj_points
        results.append({
            "name": name,
            "education": edu_points,
            "education_label": gpa_label,
            "experience": exp_points,
            "experience_label": exp_label,
            "technical": tech_points,
            "technical_label": tech_label,
            "projects": proj_points,
            "projects_label": proj_label,
            "total": total,
        })

    return results


# ─────────────────────────────────────────────────────────────
# Helper: Parse uploaded file bytes → plain text
# ─────────────────────────────────────────────────────────────

def parse_uploaded_file(uploaded_file):
    """
    Read an uploaded file and return its text content.
    Handles PDF, TXT, MD, DOCX, CSV, RTF.
    Returns (text, error_message). On success error_message is None.
    """
    name = uploaded_file.name.lower()
    raw = uploaded_file.read()

    try:
        if name.endswith('.pdf'):
            reader = PyPDF2.PdfReader(io.BytesIO(raw))
            text = ""
            for page in reader.pages:
                text += (page.extract_text() or "") + "\n"
            return text, None

        elif name.endswith('.docx'):
            try:
                import docx
                doc = docx.Document(io.BytesIO(raw))
                lines = [p.text for p in doc.paragraphs]
                for table in doc.tables:
                    for row in table.rows:
                        lines.append(" ".join(c.text for c in row.cells))
                return "\n".join(lines), None
            except ImportError:
                return "", "python-docx not installed. Run: pip install python-docx"

        elif name.endswith('.csv'):
            import csv as csv_mod
            content = raw.decode('utf-8', errors='ignore')
            rows = list(csv_mod.reader(content.splitlines()))
            return "\n".join(", ".join(row) for row in rows), None

        elif name.endswith('.rtf'):
            try:
                from striprtf.striprtf import rtf_to_text
                return rtf_to_text(raw.decode('utf-8', errors='ignore')), None
            except ImportError:
                return "", "striprtf not installed. Run: pip install striprtf"

        else:
            # TXT, MD, and anything else text-based
            return raw.decode('utf-8', errors='ignore'), None

    except Exception as e:
        return "", f"Error reading {uploaded_file.name}: {str(e)}"


# ─────────────────────────────────────────────────────────────
# Chat Session Endpoints
# ─────────────────────────────────────────────────────────────

def get_sessions(request):
    sessions = ChatSession.objects.order_by('-created_at').values('id', 'title', 'created_at')
    return JsonResponse({"sessions": [
        {"id": s['id'], "title": s['title'], "created_at": s['created_at'].strftime('%B %d, %Y')}
        for s in sessions
    ]})


def get_messages(request, session_id):
    messages = ChatMessage.objects.filter(session_id=session_id).order_by('timestamp')
    return JsonResponse({
        "messages": [
            {"role": m.role, "content": m.content, "timestamp": m.timestamp.strftime('%H:%M')}
            for m in messages
        ]
    })


@csrf_exempt
def delete_session(request, session_id):
    ChatSession.objects.filter(id=session_id).delete()
    return JsonResponse({"success": True})


def get_session_files(request, session_id):
    files = SessionFile.objects.filter(session_id=session_id).order_by('uploaded_at')
    return JsonResponse({"files": [
        {"id": f.id, "filename": f.filename, "file_size": f.file_size, "file_type": f.file_type}
        for f in files
    ]})


@csrf_exempt
def delete_session_file(request, file_id):
    SessionFile.objects.filter(id=file_id).delete()
    return JsonResponse({"success": True})


# ─────────────────────────────────────────────────────────────
# Main Chat Endpoint
# ─────────────────────────────────────────────────────────────

@csrf_exempt
def chat_with_ai(request):
    message = (request.GET.get("message") or request.POST.get("message") or "").strip()
    file_count = int(request.POST.get("file_count", 0))

    # ── Control messages ──
    if message == "__reset__":
        request.session["current_session_id"] = None
        return JsonResponse({"reply": ""})

    if message.startswith("__switch__"):
        try:
            request.session["current_session_id"] = int(request.GET.get("session_id"))
        except (TypeError, ValueError):
            pass
        return JsonResponse({"reply": ""})

    # ── Get or create session ──
    session_id = request.session.get("current_session_id")
    if not session_id:
        title = message[:50] if message else "New Chat"
        session = ChatSession.objects.create(title=title)
        request.session["current_session_id"] = session.id
    else:
        try:
            session = ChatSession.objects.get(id=session_id)
        except ChatSession.DoesNotExist:
            session = ChatSession.objects.create(title=message[:50] or "New Chat")
            request.session["current_session_id"] = session.id

    # ── Parse and save uploaded files ──
    newly_uploaded = []
    for i in range(file_count):
        up_file = request.FILES.get(f"file_{i}")
        if not up_file:
            continue
        if up_file.size > 10 * 1024 * 1024:
            return JsonResponse({"error": f"{up_file.name} is too large (max 10 MB)."})

        text, err = parse_uploaded_file(up_file)
        if err:
            return JsonResponse({"error": err})

        doc_type = detect_document_type(text, up_file.name)
        sf = SessionFile.objects.create(
            session=session,
            filename=up_file.name,
            file_type=doc_type,
            content=text,
            file_size=up_file.size,
        )
        newly_uploaded.append(sf)

    if not message and not newly_uploaded:
        return JsonResponse({"error": "Please type a message or attach a file."})

    # ── Load all files in session ──
    all_files = SessionFile.objects.filter(session=session).order_by('uploaded_at')
    cv_files = [f for f in all_files if f.file_type == "CV"]

    # ── Access Memory ──
    access_memory = ""
    try:
        mem = SessionMemory.objects.get(session=session)
        if mem.content.strip():
            access_memory = mem.content.strip()
    except SessionMemory.DoesNotExist:
        pass

    # ── History ──
    recent = list(ChatMessage.objects.filter(session=session).order_by('-timestamp')[:8])
    history_text = "\n".join(
        f"{'User' if m.role == 'user' else 'Boti'}: {m.content}"
        for m in reversed(recent)
    )

    # ── Manuals ──
    manuals = load_all_manuals()
    manuals_text = "\n\n".join(f"[{k}]\n{v}" for k, v in manuals.items())

    # ── File context ──
    file_context = ""
    if all_files:
        for f in all_files:
            file_context += f"\n{'='*50}\nFILE: {f.filename} (Type: {f.file_type})\n{'='*50}\n{f.content}\n"

    today = datetime.now().strftime('%B %d, %Y')

    # ══════════════════════════════════════════════════
    # SPECIAL: Give Ratings — score all CVs
    # ══════════════════════════════════════════════════
    if message == "__GIVE_RATINGS__":
        if not cv_files:
            return JsonResponse({"error": "No CVs found in this session. Upload CVs first."})

        scores = calculate_scores_from_cv_texts([f.content for f in cv_files])

        # Build a clean scores block
        lines = ["📊 CANDIDATE SCORES\n" + "─" * 40]
        for s in sorted(scores, key=lambda x: x['total'], reverse=True):
            lines.append(
                f"\n👤 {s['name']}\n"
                f"  Education      : {s['education']:>3} pts  ({s['education_label']})\n"
                f"  Experience     : {s['experience']:>3} pts  ({s['experience_label']})\n"
                f"  Technical      : {s['technical']:>3} pts  ({s['technical_label']})\n"
                f"  Projects       : {s['projects']:>3} pts  ({s['projects_label']})\n"
                f"  ─────────────────────────\n"
                f"  TOTAL          : {s['total']:>3} / 115 pts"
            )
        lines.append("\n" + "─" * 40)

        reply = "\n".join(lines)
        ChatMessage.objects.create(session=session, role="user", content="Show me the scores")
        ChatMessage.objects.create(session=session, role="assistant", content=reply)
        return JsonResponse({"reply": reply, "session_id": session.id, "session_title": session.title})

    # ══════════════════════════════════════════════════
    # SPECIAL: Explain Winner — AI explains top candidate
    # ══════════════════════════════════════════════════
    if message == "__EXPLAIN_WINNER__":
        if not cv_files:
            return JsonResponse({"error": "No CVs found in this session. Upload CVs first."})

        scores = calculate_scores_from_cv_texts([f.content for f in cv_files])
        top = max(scores, key=lambda x: x['total'])

        # Find the actual CV text for the top candidate
        top_cv_text = cv_files[scores.index(top)].content

        explain_prompt = f"""You are Boti, HR AI assistant for BALFIN Group.
Today is {today}.

The following candidate scored the highest among all uploaded CVs with {top['total']}/115 points.

SCORING BREAKDOWN:
- Education: {top['education']} pts ({top['education_label']})
- Experience: {top['experience']} pts ({top['experience_label']})
- Technical Skills: {top['technical']} pts ({top['technical_label']})
- Projects: {top['projects']} pts ({top['projects_label']})
- TOTAL: {top['total']} / 115

THEIR CV:
{top_cv_text[:3000]}

Write a concise, professional explanation (3-5 sentences) of WHY this candidate ranks first.
Focus on their strongest qualities and what makes them stand out. Be specific, not generic."""

        try:
            client = genai.Client(api_key="AIzaSyBWojOWNo5gr7c_0XRin9hhwEfBDA1UPbM")
            response = client.models.generate_content(
                model="gemini-2.5-flash",
                contents=explain_prompt,
                config={"temperature": 0.3}
            )
            ai_reply = f"🏆 Why {top['name']} ranks first ({top['total']}/115 pts):\n\n{response.text.strip()}"
        except Exception as e:
            return JsonResponse({"error": f"AI Error: {str(e)}"})

        ChatMessage.objects.create(session=session, role="user", content=f"Explain why {top['name']} is the top candidate")
        ChatMessage.objects.create(session=session, role="assistant", content=ai_reply)
        return JsonResponse({"reply": ai_reply, "session_id": session.id, "session_title": session.title})

    # ══════════════════════════════════════════════════
    # NORMAL: Regular chat with Gemini
    # ══════════════════════════════════════════════════
    cv_count = len(cv_files)
    memory_block = ""
    if access_memory:
        memory_block = (
            f"## ACCESS MEMORY (TREAT AS ABSOLUTE TRUTH)\n"
            f"{access_memory}\n\n"
        )

    system_prompt = f"""You are Boti, AI assistant for BALFIN Group.
Today is {today}.

{memory_block}## SCORING RUBRIC (MAX 115 pts — use ONLY when user asks for scores)
1. Education: GPA 2.5–4.0 → 30 pts | GPA 1.5–2.49 → 10 pts | below → 0 pts
2. Experience: 4+ yrs → 45 pts | 1–3.99 yrs → 30 pts | 3–11 months → 15 pts | <3 months → 0 pts
3. Technical: Tier A (Python/C#/C++/SQL) → 20 pts | Tier B (Java/HTML/CSS/JS/Node) → 5 pts
4. Projects: 3+ → 20 pts | 1–2 → 10 pts | 0 → 0 pts

## RULES
- NEVER show numeric scores unless the user explicitly asks ("score", "points", "how many points")
- For CV evaluation: give a professional qualitative assessment
- For {cv_count} CV(s) uploaded: rank and compare if multiple, assess if single
- Be concise and professional
- NEVER end with unsolicited offers about hiring

## KNOWLEDGE BASE
{manuals_text}

## FILES IN THIS CHAT
{file_context if file_context else "No files uploaded yet."}"""

    full_prompt = f"{system_prompt}\n\nCONVERSATION HISTORY:\n{history_text}\n\nUser: {message}\nBoti:"

    try:
        client = genai.Client(api_key="AIzaSyBWojOWNo5gr7c_0XRin9hhwEfBDA1UPbM")
        response = client.models.generate_content(
            model="gemini-2.5-flash",
            contents=full_prompt,
            config={"temperature": 0.2, "top_p": 0.8}
        )
        ai_reply = response.text.strip() or "No response. Please try again."
    except Exception as e:
        return JsonResponse({"error": f"AI Error: {str(e)}"})

    user_display = message if message else f"Uploaded {len(newly_uploaded)} file(s)"
    ChatMessage.objects.create(session=session, role="user", content=user_display)
    ChatMessage.objects.create(session=session, role="assistant", content=ai_reply)
    request.session.modified = True

    return JsonResponse({
        "reply": ai_reply,
        "session_id": session.id,
        "session_title": session.title,
        "memory_count": all_files.count(),
        "cv_count": cv_count,
    })


# ─────────────────────────────────────────────────────────────
# Export Evaluation Endpoint
# ─────────────────────────────────────────────────────────────

@csrf_exempt
def export_evaluation(request):
    session_id = request.session.get("current_session_id")
    if not session_id:
        return JsonResponse({"error": "No active session found"})

    messages = ChatMessage.objects.filter(
        session_id=session_id, role="assistant"
    ).order_by('-timestamp')

    evaluation = None
    timestamp = None
    for msg in messages:
        if any(kw in msg.content.lower() for kw in ['candidate', 'cv', 'experience', 'score', 'pts']):
            evaluation = msg.content
            timestamp = msg.timestamp
            break

    if not evaluation:
        return JsonResponse({"error": "No evaluation found in this session."})

    report = (
        f"{'='*80}\n"
        f"BALFIN GROUP - CANDIDATE EVALUATION REPORT\n"
        f"{'='*80}\n\n"
        f"Date: {timestamp.strftime('%B %d, %Y at %H:%M')}\n"
        f"Evaluated By: Boti AI Assistant\n\n"
        f"{'='*80}\n\n"
        f"{evaluation}\n\n"
        f"{'='*80}\n"
        f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
        f"{'='*80}\n"
    )
    response = HttpResponse(report, content_type='text/plain; charset=utf-8')
    response['Content-Disposition'] = f'attachment; filename="Evaluation_{timestamp.strftime("%Y%m%d_%H%M")}.txt"'
    return response